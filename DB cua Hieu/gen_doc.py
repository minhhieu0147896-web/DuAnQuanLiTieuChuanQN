# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

doc = Document()

# ---- Cấu hình font Times New Roman cho toàn bộ văn bản ----
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
style.paragraph_format.space_after = Pt(6)
for i in [1, 2, 3]:
    hs = doc.styles['Heading %d' % i]
    hs.font.name = 'Times New Roman'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    hs.font.color.rgb = RGBColor(33, 48, 64)

# ---- Helper functions ----
def H(lv, t):
    h = doc.add_heading(t, level=lv)
    for r in h.runs: r.font.name = 'Times New Roman'

def P(t):
    p = doc.add_paragraph(t)
    for r in p.runs: r.font.name = 'Times New Roman'

def Bul(t):
    p = doc.add_paragraph(t, style='List Bullet')
    for r in p.runs: r.font.name = 'Times New Roman'

def Code(t):
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.name = 'Consolas'; r.font.size = Pt(10)

def Sep():
    doc.add_paragraph('_' * 60)

# ============================================================
# NỘI DUNG TÀI LIỆU
# ============================================================

H(0, 'TÀI LIỆU HƯỚNG DẪN FORM LẬP THỰC ĐƠN TUẦN')
P('frmLapthucdon2 – Dành cho sinh viên năm nhất mới làm quen C#')
P('Dự án: Quản Lý Tiêu Chuẩn Quân Nhân')
Sep()

# === CHƯƠNG 1 ===
H(1, 'CHƯƠNG 1: GIỚI THIỆU TỔNG QUAN')
H(2, '1.1 Bối cảnh dự án')
P('Dự án "Quản Lý Tiêu Chuẩn Quân Nhân" là một phần mềm quản lý bếp ăn tập thể trong môi trường quân đội. Phần mềm giúp nhân viên nhà bếp lập thực đơn hàng tuần cho các chế độ ăn khác nhau (Học viên, Sĩ quan, Chiến sĩ...), theo dõi dinh dưỡng, quản lý thực phẩm và báo cáo quân số ăn hàng ngày.')
P('Form frmLapthucdon2 là form QUAN TRỌNG NHẤT trong phân hệ nhân viên nhà bếp. Đây là nơi người dùng trực tiếp chọn món ăn cho từng buổi trong tuần, xem tiến độ dinh dưỡng, và lưu thực đơn vào cơ sở dữ liệu.')

H(2, '1.2 Mục đích của form')
P('Form frmLapthucdon2 cho phép người dùng:')
Bul('Chọn tuần cần lập thực đơn thông qua DateTimePicker.')
Bul('Chọn chế độ ăn (Học viên, Sĩ quan...) từ ComboBox.')
Bul('Chọn món cho từng ô (slot) trong lưới 7 ngày × 3 buổi.')
Bul('Xem tiến độ dinh dưỡng (Đạm, Chất xơ, Chất béo) theo tuần qua ProgressBar.')
Bul('Điền thực đơn từ mẫu có sẵn.')
Bul('Lưu thực đơn xuống cơ sở dữ liệu với trạng thái Chờ duyệt.')
Bul('Xem chi tiết thống kê dinh dưỡng từng ngày bằng cách click vào tiêu đề ngày.')

H(2, '1.3 Giao diện tổng quan')
P('Giao diện form được chia thành 3 vùng chính:')
Bul('Header (màu xanh đậm): Tiêu đề "LẬP THỰC ĐƠN TUẦN" + nút Quay lại.')
Bul('Toolbar (màu trắng): DateTimePicker chọn tuần, ComboBox chọn chế độ ăn, nhãn hiển thị khoảng tuần, nút "Điền từ mẫu" và nút "Lưu thực đơn tuần".')
Bul('Body (SplitContainer chia đôi): Bên trái là lưới TableLayoutPanel 8 cột × 4 dòng. Cột đầu là tên buổi (Sáng/Trưa/Tối), 7 cột còn lại là các ngày (T2 → CN). Mỗi ô chứa 1 UserControl với các Button slot. Bên phải là Panel chọn món + 3 thanh ProgressBar dinh dưỡng + dòng trạng thái.')

Sep()

# === CHƯƠNG 2 ===
H(1, 'CHƯƠNG 2: KIẾN THỨC NỀN TẢNG CẦN NẮM')
P('Trước khi đi sâu vào code của frmLapthucdon2, bạn cần nắm vững một số khái niệm cơ bản trong C# được sử dụng xuyên suốt form này.')

H(2, '2.1 Dictionary<K, V> – Từ điển (Bảng băm)')
P('Dictionary là một cấu trúc dữ liệu lưu trữ cặp Key-Value (Khóa-Giá trị). Hãy tưởng tượng nó giống như một cuốn từ điển giấy: bạn tra từ "Apple" (Key) thì nhận được nghĩa "Quả táo" (Value). Mỗi Key là duy nhất. Tra cực nhanh với độ phức tạp O(1) — không cần duyệt tuần tự.')
P('Cú pháp cơ bản:')
Code('Dictionary<string, MonAnModel> dict = new Dictionary<string, MonAnModel>();')
Code('dict["key1"] = monAn;            // Thêm hoặc cập nhật')
Code('MonAnModel m = dict["key1"];     // Lấy ra, lỗi nếu key không tồn tại')
Code('dict.ContainsKey("key1");         // Kiểm tra key có tồn tại không')
Code('dict.TryGetValue("key1", out m); // Lấy an toàn, không gây lỗi')
Code('dict.Remove("key1");             // Xóa')
P('Trong frmLapthucdon2, Dictionary được dùng RẤT NHIỀU vì form cần ánh xạ nhanh từ mã định danh ô (key chuỗi) → món ăn (MonAnModel) hoặc → nút Button.')

H(2, '2.2 Enum – Kiểu liệt kê')
P('Enum (Enumeration) là kiểu dữ liệu đặc biệt cho phép bạn định nghĩa một tập hợp các hằng số có tên. Thay vì dùng số 0, 1, 2 khó hiểu, bạn dùng tên có ý nghĩa như DishCategory.Mặn, MealKind.Sáng.')
Code('public enum DishCategory     // Các loại món ăn')
Code('{')
Code('    Man,           // Món mặn')
Code('    Canh,          // Canh / súp')
Code('    Rau,           // Rau xào / luộc')
Code('    TrangMieng,    // Tráng miệng / trái cây')
Code('    SuaHop,        // Sữa hộp (chỉ cho Học viên)')
Code('    Com            // Cơm trắng (tự động mọi buổi)')
Code('}')
Code('public enum MealKind { Sang = 0, Trua = 1, Toi = 2 }')
P('Khi dùng enum, code trở nên rõ ràng: DishCategory.Canh dễ đọc hơn số 1. Enum còn giúp trình biên dịch bắt lỗi nếu bạn gõ sai tên.')

H(2, '2.3 UserControl – Điều khiển tùy chỉnh')
P('UserControl trong WinForms giống như một "form con". Bạn có thể thiết kế giao diện bằng cách kéo thả các control (Button, Label, Panel...) vào nó, rồi tái sử dụng UserControl đó ở nhiều nơi khác nhau.')
P('Trong frmLapthucdon2, mỗi ô trong lưới 7 ngày × 3 buổi là một UserControl:')
Bul('UcMealCellSang: Dùng cho buổi Sáng. Chứa 4 nút: Mặn, Canh, Sữa hộp, Cơm.')
Bul('UcMealCellTruaToi: Dùng cho buổi Trưa và buổi Tối. Chứa 8 nút: 4 Mặn, Canh, Rau, Tráng miệng, Cơm.')
P('Nhờ UserControl, form chính không cần tạo 63 Button riêng lẻ — chỉ cần 21 UserControl, mỗi cái tự quản lý các nút bên trong.')

H(2, '2.4 Sự kiện (Event) trong WinForms')
P('WinForms hoạt động theo mô hình "hướng sự kiện" (event-driven). Thay vì chạy tuần tự từ trên xuống dưới, chương trình chờ người dùng tương tác (click, gõ phím, chọn ngày...) rồi mới thực thi hàm xử lý tương ứng.')
P('Ví dụ: Tất cả Button slot đều được gán chung MỘT hàm xử lý Slot_Click:')
Code('btn.Click += Slot_Click;')
Code('private void Slot_Click(object sender, EventArgs e)')
Code('{')
Code('    Button clickedBtn = sender as Button; // sender = nút bị click')
Code('    SlotMetadata meta = clickedBtn.Tag as SlotMetadata; // đọc metadata')
Code('    // ... xử lý tiếp ...')
Code('}')

H(2, '2.5 string.Format – Định dạng chuỗi')
P('string.Format() là cách tạo chuỗi có chứa giá trị biến mà không cần cộng chuỗi thủ công. Dùng placeholder {0}, {1}, {2}... để đánh dấu vị trí cần chèn. Dấu hai chấm dùng cho format specifier (định dạng đặc biệt).')
Code('// Ví dụ tạo Key cho một slot:')
Code('string key = string.Format("{0:yyyyMMdd}-{1}-{2}-{3}", date, buoiAnId, category, index);')
Code('// Kết quả: "20260609-1-Man-0"')
P('{0:yyyyMMdd} nghĩa là: lấy giá trị thứ 0 (date), định dạng thành chuỗi 8 chữ số yyyyMMdd.')

Sep()

# === CHƯƠNG 3 ===
H(1, 'CHƯƠNG 3: KIẾN TRÚC DỮ LIỆU CỦA FORM')
P('Đây là chương QUAN TRỌNG NHẤT. Form quản lý ~63 ô (7 ngày × 3 buổi × ~3 loại món), mỗi ô chứa 1 món ăn. Để quản lý hiệu quả, form dùng một hệ thống Key thống nhất và 4 bộ đệm (buffer) phối hợp với nhau.')

H(2, '3.1 Hệ thống Key – "Chứng minh thư" của mỗi ô')
P('Mỗi ô (slot) trong lưới được định danh bằng MỘT CHUỖI KEY DUY NHẤT. Key này được dùng làm khóa trong tất cả Dictionary của form.')
P('Cấu trúc Key:')
Code('"yyyyMMdd" + "-" + buoiAnId + "-" + DishCategory + "-" + categoryIndex')
P('Ví dụ cụ thể:')
Code('"20260609-1-Man-0"  →  Ngày 09/06/2026, buổi Sáng (id=1), món Mặn, vị trí thứ 1')
Code('"20260610-2-Man-3"  →  Ngày 10/06/2026, buổi Trưa (id=2), món Mặn thứ 4')
Code('"20260610-3-Com-0"  →  Ngày 10/06/2026, buổi Tối (id=3), Cơm trắng')
P('Tại sao cần Key? Vì form cần biết chính xác món "Thịt kho tàu" đang nằm ở ô nào trong 63 ô. Với Dictionary + Key, tìm kiếm trong O(1) — gần như tức thì. Không có Key, mỗi lần tìm phải duyệt qua tất cả 63 Button.')

H(2, '3.2 Bốn bộ đệm (buffer) chính')
P('Form sử dụng 4 cấu trúc dữ liệu chính. Tất cả đều dùng chung hệ thống Key ở trên. Hãy hiểu chúng như 4 "hộp đựng" khác nhau, mỗi hộp có một mục đích riêng:')

H(3, '3.2.1 _slots – "Danh bạ" các nút trên giao diện')
Code('Dictionary<string, Button> _slots')
P('Ánh xạ từ Key → Button (nút bấm trên màn hình). Khi cần tô màu, đổi text, khóa/mở khóa một ô, chỉ cần gọi _slots["key"] là có ngay Button cần thao tác. Giống như cuốn danh bạ: biết số điện thoại (Key) → tìm được ngay người cần gọi (Button).')

H(3, '3.2.2 _selectedMeals – "Ai đang ngồi ở ô nào"')
Code('Dictionary<string, MonAnModel> _selectedMeals')
P('Ánh xạ từ Key → Món ăn hiện được chọn. Đây là BỘ NHỚ LÀM VIỆC của form — mọi thao tác chọn món, xóa món đều ghi vào đây. Khi nhấn nút Lưu, toàn bộ nội dung _selectedMeals được ghi xuống database. Khi cần vẽ lại giao diện (RefreshSlots), form duyệt _selectedMeals để biết ô nào có món gì mà tô màu/text.')

H(3, '3.2.3 _weeklyBuffers – "Nhật ký các tuần đã xem"')
Code('Dictionary<int, Dictionary<string, MonAnModel>> _weeklyBuffers')
P('Đây là Dictionary LỒNG Dictionary — cấu trúc hơi khó hiểu nếu mới học. Hãy tách ra:')
Bul('Lớp NGOÀI: Key = int (mã tuần, vd: 20260601). Value = một Dictionary khác.')
Bul('Lớp TRONG: Key = string (key của slot). Value = MonAnModel (món ăn).')
P('Mục đích: khi người dùng chuyển qua lại giữa các tuần (xem tuần 1 → tuần 2 → quay lại tuần 1), form KHÔNG query lại database. Nó lấy dữ liệu từ _weeklyBuffers ra, tiết kiệm thời gian đáng kể.')
Code('// Ví dụ cách dùng:')
Code('_weeklyBuffers[20260601]  →  Dictionary chứa toàn bộ món của tuần 01/06/2026')
Code('// Khi quay lại tuần cũ, chỉ cần:')
Code('_selectedMeals = _weeklyBuffers[20260601];  // Không query DB!')

H(3, '3.2.4 _weeklyStatuses – Trạng thái từng tuần')
Code('Dictionary<int, string> _weeklyStatuses')
P('Lưu trạng thái của mỗi tuần, đi kèm với _weeklyBuffers:')
Bul('"NhapLieu" – Đang soạn thảo, có thể sửa thoải mái.')
Bul('"ChoDuyet" – Đã gửi lên cấp trên chờ duyệt. Vẫn sửa được nhưng sẽ tự động chuyển về NhapLieu.')
Bul('"DaDuyet" – Đã được duyệt chính thức. Form bị KHÓA TOÀN BỘ, chỉ được xem.')

H(2, '3.3 Mối quan hệ giữa các bộ đệm & Sơ đồ luồng dữ liệu')
P('MỐI QUAN HỆ CỐT LÕI:')
Bul('_slots["key"] trả lời câu hỏi: "Nút bấm của ô này NẰM Ở ĐÂU trên màn hình?"')
Bul('_selectedMeals["key"] trả lời câu hỏi: "Ô này đang CHỨA MÓN GÌ?"')
P('HAI Dictionary này DÙNG CHUNG KEY. Một cái quản lý GIAO DIỆN, một cái quản lý DỮ LIỆU.')
P('SƠ ĐỒ LUỒNG DỮ LIỆU:')
Code('DATABASE (SQL Server)')
Code('    ↓ LoadWeek() query / btnluu_Click() INSERT')
Code('_weeklyBuffers[maTuan]  ←── Bộ đệm, tránh query lại ──┐')
Code('    ↓ copy sang (nếu chưa có trong cache)              │')
Code('_selectedMeals (bộ nhớ làm việc HIỆN TẠI) ─────────────┘')
Code('    ↓ RefreshSlots() duyệt _selectedMeals')
Code('_slots (Button trên màn hình)')

Sep()

# === CHƯƠNG 4 ===
H(1, 'CHƯƠNG 4: CHI TIẾT LOGIC HOẠT ĐỘNG')

H(2, '4.1 Khởi tạo form – Constructor & Load')
P('Khi form được mở, hai hàm sau chạy theo thứ tự:')
P('CONSTRUCTOR (hàm tạo):')
Bul('InitializeComponent() – Hàm do Designer tự sinh, tạo toàn bộ control trên form.')
Bul('Khởi tạo ToolTip, bật DoubleBuffered (chống giật màn hình).')
Bul('Gọi InitGridReferenceArrays() – tạo mảng tham chiếu nhanh _dayHeaders (7 Label tiêu đề ngày) và _mealCells (UserControl[3,7]).')
Bul('Đăng ký sự kiện Load và FormClosing.')
P('SỰ KIỆN LOAD (frmLapthucdon2_Load) – chạy MỘT LẦN DUY NHẤT nhờ cờ _isInitialized:')
Bul('WireEvents() – Gán sự kiện Click cho tất cả nút: btnLuu, btnDienTuMau, btnXoaMon, btnReload, dtpWeek, cboCheDo...')
Bul('SetupFormAppearance() – Chỉnh font Segoe UI, màu nền, style các nút.')
Bul('LoadReferenceData() – Load danh sách buổi ăn (Sáng/Trưa/Tối) và danh sách chế độ từ DB.')
Bul('LoadWeek() – Load dữ liệu thực đơn của tuần hiện tại.')

H(2, '4.2 Load dữ liệu tham chiếu – LoadReferenceData()')
P('Hàm này chạy MỘT LẦN khi form mở, load hai thứ:')
Bul('Danh sách buổi ăn (_buoiAns): Lấy từ bảng Buoi_an qua BuoiAnDAO.GetAll(). Thường gồm 3 buổi: Sáng (id=1), Trưa (id=2), Tối (id=3).')
Bul('Danh sách chế độ ăn: Đổ vào ComboBox cboCheDo từ B_QN.GetAllCheDo().')
P('Cờ _isLoadingReferenceData = true được bật trong lúc load để ngăn sự kiện SelectedIndexChanged của ComboBox vô tình kích hoạt LoadWeek() khi dữ liệu chưa sẵn sàng. Đây là kỹ thuật "dùng cờ khóa tạm" rất phổ biến trong WinForms.')

H(2, '4.3 LoadWeek() – "Trái tim" của form')
P('LoadWeek() là hàm QUAN TRỌNG NHẤT. Nó chạy mỗi khi người dùng đổi tuần hoặc đổi chế độ ăn. Trình tự từng bước:')
P('BƯỚC 1 – Tính weekKey: Từ ngày được chọn trên DateTimePicker, tính ra ngày Thứ 2 của tuần, rồi chuyển thành mã int: yyyyMMdd (vd: 20260601).')
P('BƯỚC 2 – Kiểm tra bộ đệm:')
Bul('Nếu _weeklyBuffers đã có weekKey → copy sang _selectedMeals. KHÔNG query DB.')
Bul('Nếu chưa có → gọi LoadExistingMealsFromDatabase() để query DB, rồi lưu kết quả vào _weeklyBuffers[weekKey].')
P('BƯỚC 3 – Cập nhật giao diện:')
Bul('Hiển thị lblWeekRange: "Tuần 01/06/2026 - 07/06/2026".')
Bul('Lấy chỉ tiêu dinh dưỡng tuần từ MonAnDAO.GetWeeklyNutritionTarget().')
Bul('PopulateWeekGrid() – Gán metadata (SlotMetadata) vào Button.Tag cho từng nút.')
Bul('ApplyAutomaticMilkSlots() – Tự động điền Sữa hộp nếu là chế độ Học viên.')
Bul('ApplyAutomaticRiceSlots() – Tự động điền Cơm trắng cho mọi chế độ, mọi buổi.')
Bul('RefreshSlots() – Vẽ lại text & màu nền cho tất cả Button.')
Bul('UpdateStatus() – Cập nhật 3 thanh ProgressBar & dòng trạng thái.')
Bul('ApplyEditPermissions() – Khóa form nếu tuần đã được duyệt.')

H(2, '4.4 PopulateWeekGrid() – Gán metadata cho Button')
P('Hàm này KHÔNG tạo mới control! Tất cả UserControl đã được tạo sẵn trong Designer từ lúc thiết kế form. PopulateWeekGrid chỉ làm 2 việc:')
P('VIỆC 1 – Nhét metadata vào Button.Tag: Mỗi Button có property Tag (kiểu object). Form tạo một đối tượng SlotMetadata chứa: Key, Date, BuoiAn, MealKind, DishCategory, CategoryIndex — rồi gán vào Tag. Sau này khi click vào Button, chỉ cần đọc Tag là biết ngay ô đó thuộc ngày nào, buổi nào.')
P('VIỆC 2 – Đăng ký sự kiện: Tất cả Button đều được gán chung MỘT hàm Slot_Click. Khi người dùng click bất kỳ ô nào, Slot_Click được gọi, đọc Tag để xác định ô nào đang được chọn.')

H(2, '4.5 Slot_Click() – Xử lý khi người dùng click vào ô')
P('Trình tự xử lý khi click vào một Button slot:')
Bul('B1 – Lấy SlotMetadata từ clickedBtn.Tag. Nếu null → return.')
Bul('B2 – Kiểm tra quyền: nếu _currentWeekStatus là "DaDuyet" → hiện MessageBox "Thực đơn đã duyệt, không thể sửa" → return.')
Bul('B3 – Tô viền: bỏ BorderSize=1 cho tất cả slot, gán BorderSize=3 cho slot đang chọn. Đây là cách đánh dấu trực quan.')
Bul('B4 – Hiển thị thông tin: lblActiveSlot = "09/06 - Sáng\nMặn 1".')
Bul('B5 – Xử lý đặc biệt: nếu là slot Sữa hộp hoặc Cơm trắng → return ngay, KHÔNG mở form chọn món (vì đây là món tự động, không được phép đổi).')
Bul('B6 – Mở frmchonmon: form danh sách món ăn được lọc theo loại (Mặn/Canh/Rau/Tráng miệng). Nếu chọn OK → lưu vào _selectedMeals[meta.Key] → RefreshSlots() → UpdateStatus().')

H(2, '4.6 Tự động điền Sữa hộp và Cơm trắng')
P('Hai món được hệ thống TỰ ĐỘNG thêm, người dùng KHÔNG cần và KHÔNG THỂ chọn:')
P('SỮA HỘP (ApplyAutomaticMilkSlots):')
Bul('Chỉ áp dụng khi chế độ là Học viên (IsHocVienSelected() == true).')
Bul('Gọi MonAnDAO.Instance.GetOrCreateSua(). Nếu DB chưa có món "Sữa", tự động INSERT với dinh dưỡng: Đạm 6g, Chất béo 6g, Chất xơ 0g.')
Bul('Duyệt tất cả slot trong _slots. Nếu DishCategory == SuaHop → _selectedMeals[key] = milk.')
P('CƠM TRẮNG (ApplyAutomaticRiceSlots):')
Bul('Áp dụng cho TẤT CẢ chế độ, TẤT CẢ buổi (Sáng/Trưa/Tối).')
Bul('Gọi MonAnDAO.Instance.GetOrCreateCom(). Nếu DB chưa có, tự tạo món "Cơm trắng" với dinh dưỡng: Đạm 4g, Chất béo 0.5g, Chất xơ 0.5g (~1 suất cơm 200g).')
Bul('Duyệt tất cả slot. Nếu DishCategory == Com → _selectedMeals[key] = com.')

H(2, '4.7 btnluu_Click() – Lưu thực đơn tuần xuống database')
P('Quy trình 6 bước khi nhấn nút "Lưu thực đơn tuần":')
Bul('B1 – ValidateWeek(): Duyệt TẤT CẢ slot trong _slots. Slot nào chưa có món trong _selectedMeals → báo lỗi "Còn thiếu món: 09/06 - Sáng - Mặn 1" → return, không cho lưu.')
Bul('B2 – MessageBox xác nhận: "Lưu toàn bộ thực đơn trong tuần này vào cơ sở dữ liệu?"')
Bul('B3 – XÓA dữ liệu cũ: Với mỗi ngày trong tuần, gọi ThucDonDAO.GetOrCreate() để lấy/cấp bản ghi Thuc_don. Sau đó gọi ChiTietThucDonDAO.DeleteByThucDonNgayBuoi() để xóa toàn bộ chi tiết món cũ. (Xóa trước, chèn sau để tránh trùng lặp.)')
Bul('B4 – INSERT từng món: Duyệt _slots, slot nào có món trong _selectedMeals → ChiTietThucDonDAO.Insert(thucDonId, date, buoiAnId, monAnId).')
Bul('B5 – Cập nhật trạng thái: ThucDonDAO.UpdateTrangThai(thucDonId, "ChoDuyet").')
Bul('B6 – Dọn dẹp: Xóa _weeklyBuffers[weekKey] và _weeklyStatuses[weekKey] (để lần sau load mới từ DB). Gọi LoadWeek() để hiển thị trạng thái mới.')

Sep()

# === CHƯƠNG 5 ===
H(1, 'CHƯƠNG 5: GIẢI THÍCH TỪNG HÀM QUAN TRỌNG')

H(2, '5.1 GetRequiredCategories(MealKind meal)')
P('Trả về danh sách các loại món CẦN CÓ cho một buổi, dùng yield return (iterator pattern):')
Bul('Sáng: Mặn, Canh (+ Sữa nếu Học viên), Cơm.')
Bul('Trưa/Tối: 4 Mặn, Canh, Rau, Tráng miệng, Cơm.')

H(2, '5.2 GetCategoryCount(MealKind meal, DishCategory category)')
P('Trả về SỐ LƯỢNG slot cho một loại món trong một buổi:')
Bul('Sáng: tất cả loại món đều 1 slot.')
Bul('Trưa/Tối: Mặn = 4 slot, còn lại (Canh, Rau, Tráng miệng, Cơm) = 1 slot.')

H(2, '5.3 BuildKey(DateTime date, int buoiAnId, DishCategory category, int index)')
P('Tạo chuỗi Key định danh cho một slot. Đây là hàm được gọi RẤT NHIỀU LẦN trong form.')
Code('return string.Format("{0:yyyyMMdd}-{1}-{2}-{3}", date, buoiAnId, category, index);')

H(2, '5.4 RefreshSlots()')
P('Làm mới giao diện toàn bộ lưới: duyệt tất cả Button trong _slots, kiểm tra _selectedMeals:')
Bul('Ô CÓ món → hiển thị tên món (vd: "Thịt kho tàu"), nền trắng.')
Bul('Ô TRỐNG → hiển thị placeholder (vd: "+ Mặn 1"), nền màu theo loại món (vàng cho Mặn, xanh cho Canh, hồng cho Tráng miệng...).')

H(2, '5.5 UpdateStatus()')
P('Cập nhật 2 thứ:')
Bul('3 thanh ProgressBar: Tính tổng Đạm/Chất xơ/Chất béo từ Sum(_selectedMeals), so với chỉ tiêu tuần từ DB. Nếu vượt 100% → thanh chuyển ĐỎ (qua Win32 API PBM_SETSTATE).')
Bul('Dòng trạng thái lblStatus: Hiển thị "Đã chọn X/Y ô", trạng thái tuần, ràng buộc từng buổi, cảnh báo dinh dưỡng.')

H(2, '5.6 ValidateWeek()')
P('Hàm kiểm tra tính đầy đủ của thực đơn TRƯỚC KHI LƯU. Duyệt qua tất cả slot trong _slots theo thứ tự ngày → buổi. Nếu slot nào chưa có món trong _selectedMeals, trả về chuỗi lỗi (vd: "Còn thiếu món: 09/06/2026 - Sáng - Canh 1"). Nếu đầy đủ → trả về null.')

H(2, '5.7 ClassifyDishType(string value)')
P('Phân loại món ăn dựa vào chuỗi loại món (monan_loaimon) từ database. Dùng NormalizeText() để chuẩn hóa tiếng Việt (bỏ dấu, lowercase), sau đó dò từ khóa:')
Bul('Chứa "sua"/"milk" → SuaHop')
Bul('Chứa "com"/"cơm" → Com')
Bul('Chứa "canh"/"sup" → Canh')
Bul('Chứa "rau"/"xao"/"luoc" → Rau')
Bul('Chứa "trang mieng"/"trai cay" → TrangMieng')
Bul('Còn lại → Man (Món mặn)')

Sep()

# === PHỤ LỤC ===
H(1, 'PHỤ LỤC A: BẢNG THUẬT NGỮ')
P('Slot / Ô – Một ô trong lưới thực đơn, đại diện cho 1 món ăn tại 1 buổi, 1 ngày, 1 loại món.')
P('Key – Chuỗi định danh duy nhất cho mỗi slot, định dạng: yyyyMMdd-buoiAnId-DishCategory-index.')
P('Bộ đệm (Buffer/Cache) – Vùng nhớ tạm lưu dữ liệu để tránh truy vấn database nhiều lần.')
P('Metadata – Dữ liệu mô tả về một đối tượng. SlotMetadata mô tả một ô (ngày, buổi, loại món...).')
P('Tag – Property kiểu object của mọi Control trong WinForms, dùng để "gắn" dữ liệu tùy ý vào control.')
P('UserControl – Control tổng hợp do người dùng tự thiết kế, chứa nhiều control con bên trong.')
P('FlowLayoutPanel – Panel tự động sắp xếp các control con theo chiều dọc hoặc ngang.')
P('DoubleBuffered – Kỹ thuật vẽ vào bộ đệm ẩn trước rồi mới hiển thị, giúp giảm giật màn hình.')
P('yield return – Từ khóa C# cho phép trả về từng phần tử một (iterator pattern), tiết kiệm bộ nhớ.')
P('O(1) – Độ phức tạp thuật toán: thời gian thực thi không đổi, không phụ thuộc kích thước dữ liệu.')
P('DishCategory – Enum liệt kê các loại món: Man, Canh, Rau, TrangMieng, SuaHop, Com.')
P('MealKind – Enum liệt kê các buổi: Sang=0, Trua=1, Toi=2.')

Sep()
H(1, 'PHỤ LỤC B: SƠ ĐỒ QUAN HỆ GIỮA CÁC LỚP')
Code('frmLapthucdon2 (Form chính)')
Code('├── _slots: Dictionary<string, Button>')
Code('│   └── Key → Button (ô trên giao diện)')
Code('├── _selectedMeals: Dictionary<string, MonAnModel>')
Code('│   └── Key → Món ăn đang được chọn')
Code('├── _weeklyBuffers: Dictionary<int, Dictionary<string, MonAnModel>>')
Code('│   └── Mã tuần → (Key → Món ăn)  ← Cache')
Code('├── _weeklyStatuses: Dictionary<int, string>')
Code('│   └── Mã tuần → "NhapLieu"|"ChoDuyet"|"DaDuyet"')
Code('├── _mealCells: UserControl[3, 7]   ← 3 buổi × 7 ngày')
Code('│   ├── UcMealCellSang (Sáng)')
Code('│   │   └── btnMan1, btnCanh1, btnSua1, btnCom1')
Code('│   └── UcMealCellTruaToi (Trưa & Tối)')
Code('│       └── btnMan1..4, btnCanh1, btnRau1, btnTrangMieng1, btnCom1')
Code('├── SlotMetadata (lưu trong Button.Tag)')
Code('│   └── Key, Date, BuoiAn, MealKind, DishCategory, CategoryIndex')
Code('└── DAO Layer (giao tiếp database)')
Code('    ├── MonAnDAO (GetOrCreateSua, GetOrCreateCom, ...)')
Code('    ├── ThucDonDAO (GetOrCreate, UpdateTrangThai)')
Code('    ├── ChiTietThucDonDAO (GetByThucDonNgay, Insert, Delete...)')
Code('    └── BuoiAnDAO (GetAll)')

Sep()
H(1, 'LỜI KẾT')
P('Form frmLapthucdon2 là một ví dụ điển hình về cách tổ chức code trong ứng dụng WinForms thực tế. Những kiến thức bạn học được từ form này — Dictionary, Enum, UserControl, Event-driven programming, buffer/cache pattern — sẽ theo bạn trong suốt sự nghiệp lập trình, không chỉ giới hạn ở C# hay WinForms.')
P('Lời khuyên thực tế: Mở Visual Studio, đặt breakpoint (F9) ở đầu hàm LoadWeek(), chạy debug (F5), và theo dõi từng bước dữ liệu chạy qua các bộ đệm như thế nào. Đó là cách học nhanh nhất và trực quan nhất.')
P('Chúc bạn học tốt!')
P('--- HẾT ---')

# ============================================================
# LƯU FILE
# ============================================================
output = r'E:\lung tung\DuAnQuanLiTieuChuanQN\DB cua Hieu\TaiLieu_frmLapthucdon2.docx'
doc.save(output)
print('DA TAO XONG:', output)
